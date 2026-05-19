"""
选股小龙虾 - docx报告生成
生成选股结果docx报告，包含TOP20和TOP8
"""
import logging
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def generate_selection_report(stocks: list, output_path: str, title: str = "选股小龙虾") -> str:
    """
    生成选股结果docx报告
    
    Args:
        stocks: 股票列表（已评分排序）
        output_path: 输出文件路径
        title: 报告标题
    
    Returns:
        输出文件路径
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    # 标题
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"更新时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 市场概览（如果有）
    if stocks and "market" in stocks[0]:
        market = stocks[0]["market"]
        p = doc.add_paragraph()
        p.add_run(f"大盘评分：{market.get('total_score', 0):.0f}分 | {market.get('advice', '')}").italic = True
    
    doc.add_paragraph()  # 空行
    
    # TOP20表格
    doc.add_heading(f"TOP{len(stocks)} 综合评分股", level=1)
    _add_stock_table(doc, stocks, show_overnight=False)
    
    # 如果有隔夜数据，添加隔夜推荐
    overnight_stocks = [s for s in stocks if s.get("overnight_rank")]
    if overnight_stocks:
        doc.add_heading("TOP8 隔夜推荐", level=1)
        _add_stock_table(doc, overnight_stocks, show_overnight=True)
    
    # 保存
    doc.save(output_path)
    logger.info(f"报告生成完成: {output_path}")
    return output_path


def _add_stock_table(doc: Document, stocks: list, show_overnight: bool = False):
    """
    添加股票表格
    
    列：排名 | 代码 | 名称 | 收盘价 | 涨跌幅 | 量比 | 换手率 | 总分 | ...
    """
    # 表格标题行
    if show_overnight:
        headers = ["排名", "代码", "名称", "收盘价", "涨幅%", "量比", "换手率%", "隔夜总分", "K线", "套牢", "备注"]
    else:
        headers = ["排名", "代码", "名称", "收盘价", "涨幅%", "量比", "换手率%", "总分", "资金", "基本面", "技术", "风控"]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 标题行
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        for paragraph in header_row.cells[i].paragraphs:
            if paragraph.runs:
                paragraph.runs[0].bold = True
    
    # 数据行
    for stock in stocks[:20]:  # 最多20行
        row = table.add_row()
        cells = row.cells
        
        if show_overnight:
            cells[0].text = str(stock.get("overnight_rank", "-"))
            cells[1].text = stock.get("code", "-")
            cells[2].text = stock.get("name", "-")[:6]
            cells[3].text = f"{stock.get('price', 0):.2f}"
            cells[4].text = f"{stock.get('pct_change', 0):.2f}%"
            cells[5].text = f"{stock.get('volume_ratio', 0):.2f}"
            cells[6].text = f"{stock.get('turnover_rate', 0):.2f}%"
            cells[7].text = f"{stock.get('overnight_total_score', 0):.1f}"
            cells[8].text = stock.get("kline_direction", "-")[:4]
            cells[9].text = f"{stock.get('trapped_level', 0)}级"
            level = stock.get("trapped_level", 0)
            emoji = "⚠️" if level >= 4 else ""
            cells[10].text = emoji
        else:
            cells[0].text = str(stock.get("rank", "-"))
            cells[1].text = stock.get("code", "-")
            cells[2].text = stock.get("name", "-")[:6]
            cells[3].text = f"{stock.get('price', 0):.2f}"
            cells[4].text = f"{stock.get('pct_change', 0):.2f}%"
            cells[5].text = f"{stock.get('volume_ratio', 0):.2f}"
            cells[6].text = f"{stock.get('turnover_rate', 0):.2f}%"
            cells[7].text = f"{stock.get('score_total', 0):.1f}"
            cells[8].text = f"{stock.get('score_capital', 0):.0f}"
            cells[9].text = f"{stock.get('score_fundamental', 0):.0f}"
            cells[10].text = f"{stock.get('score_technical', 0):.0f}"
            if len(cells) > 11:
                cells[11].text = f"{stock.get('score_risk', 0):.0f}"
    
    # 设置列宽
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(0.8)


def generate_overnight_report(stocks: list, output_path: str) -> str:
    """
    生成隔夜套利专项报告
    
    Args:
        stocks: TOP8隔夜推荐列表
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    doc = Document()
    
    # 标题
    heading = doc.add_heading("选股小龙虾 - 隔夜套利推荐", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"更新时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 简要说明
    p1 = doc.add_paragraph("策略说明：尾盘强势股，次日高开概率高")
    p1.runs[0].italic = True
    doc.add_paragraph("评分构成：基础分(100分) + K线修正(±15分) + 套牢盘修正(±10分)")
    doc.add_paragraph()
    
    # 表格
    _add_stock_table(doc, stocks, show_overnight=True)
    
    doc.save(output_path)
    logger.info(f"隔夜报告生成: {output_path}")
    return output_path
