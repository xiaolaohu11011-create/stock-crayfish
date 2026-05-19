#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
选股小龙虾 - 持仓健康检查
每日扫描持仓，生成卖出信号和持仓报告
"""
import sys
import os
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Optional

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from position_tracker import (
    get_positions, sell_position, update_position,
    calculate_position_value, POSITION_FILE
)
from data.eastmoney import fetch_realtime_quote
from strategy.market_analysis import analyze_market
from config.config import MARKET_SCORE_THRESHOLD

logger = logging.getLogger(__name__)


# ============ 卖出信号规则 ============

class SellSignal:
    """卖出信号定义"""
    
    # 信号优先级（数字越小越优先）
    PRIORITY = {
        "market_crash": 1,      # 大盘崩盘
        "stop_loss": 2,         # 止损
        "take_profit": 3,       # 止盈
        "trapped_severe": 4,    # 严重套牢
        "score_drop": 5,        # 评分大幅下降
        "kline_down": 6,        # K线转空
        "time_stop": 7,         # 时间止损
        "capital_outflow": 8,   # 资金大幅流出
    }
    
    def __init__(self, code: str, reason: str, detail: str, priority: int, suggested_action: str):
        self.code = code
        self.reason = reason
        self.detail = detail
        self.priority = priority
        self.suggested_action = suggested_action
    
    def to_dict(self) -> Dict:
        return {
            "code": self.code,
            "reason": self.reason,
            "detail": self.detail,
            "priority": self.priority,
            "suggested_action": self.suggested_action,
        }


def check_market_condition(market: Dict) -> Optional[SellSignal]:
    """
    检查大盘环境
    
    规则：大盘评分 < 40 分，建议清仓
    """
    score = market.get("total_score", 0)
    
    if score < 40:
        return SellSignal(
            code="ALL",
            reason="market_crash",
            detail=f"大盘评分仅{score:.0f}分，市场极度弱势",
            priority=SellSignal.PRIORITY["market_crash"],
            suggested_action="清仓观望",
        )
    elif score < MARKET_SCORE_THRESHOLD:
        return SellSignal(
            code="ALL",
            reason="market_weak",
            detail=f"大盘评分{score:.0f}分，低于{MARKET_SCORE_THRESHOLD}分阈值",
            priority=99,  # 低优先级，仅提醒
            suggested_action="减仓或观望",
        )
    
    return None


def check_stop_loss(position: Dict, current_price: float) -> Optional[SellSignal]:
    """
    检查止损条件
    
    规则：当前价 <= 止损价（默认买入价-5%）
    """
    stop_loss = position.get("stop_loss", 0)
    
    if stop_loss > 0 and current_price <= stop_loss:
        loss_pct = round((current_price - position["buy_price"]) / position["buy_price"] * 100, 2)
        return SellSignal(
            code=position["code"],
            reason="stop_loss",
            detail=f"当前价{current_price} <= 止损价{stop_loss}，亏损{loss_pct}%",
            priority=SellSignal.PRIORITY["stop_loss"],
            suggested_action="立即止损卖出",
        )
    
    return None


def check_take_profit(position: Dict, current_price: float) -> Optional[SellSignal]:
    """
    检查止盈条件
    
    规则：当前价 >= 止盈价（默认买入价+10%）
    """
    take_profit = position.get("take_profit", 0)
    
    if take_profit > 0 and current_price >= take_profit:
        profit_pct = round((current_price - position["buy_price"]) / position["buy_price"] * 100, 2)
        return SellSignal(
            code=position["code"],
            reason="take_profit",
            detail=f"当前价{current_price} >= 止盈价{take_profit}，盈利{profit_pct}%",
            priority=SellSignal.PRIORITY["take_profit"],
            suggested_action="止盈卖出（可留半仓）",
        )
    
    return None


def check_time_stop(position: Dict) -> Optional[SellSignal]:
    """
    检查时间止损
    
    规则：持有超过5天且亏损
    """
    buy_date = datetime.strptime(position["buy_date"], "%Y-%m-%d")
    hold_days = (datetime.now() - buy_date).days
    
    if hold_days >= 5:
        return SellSignal(
            code=position["code"],
            reason="time_stop",
            detail=f"持有{hold_days}天，超过5天时间止损线",
            priority=SellSignal.PRIORITY["time_stop"],
            suggested_action="时间止损卖出",
        )
    
    return None


def check_score_drop(position: Dict, current_score: float) -> Optional[SellSignal]:
    """
    检查评分下降
    
    规则：当前综合评分较买入时下降超过15分
    （需要外部传入当前评分，或从选股结果获取）
    """
    buy_score = position.get("buy_score", 0)
    
    if buy_score > 0 and current_score > 0:
        drop = buy_score - current_score
        if drop >= 15:
            return SellSignal(
                code=position["code"],
                reason="score_drop",
                detail=f"评分从{buy_score}降至{current_score}，下降{drop}分",
                priority=SellSignal.PRIORITY["score_drop"],
                suggested_action="评分恶化，考虑卖出",
            )
    
    return None


def check_trapped_level(position: Dict, trapped_level: int) -> Optional[SellSignal]:
    """
    检查套牢盘恶化
    
    规则：套牢盘达到4级（较重）或5级（严重）
    """
    if trapped_level >= 5:
        return SellSignal(
            code=position["code"],
            reason="trapped_severe",
            detail=f"套牢盘达到{trapped_level}级（严重套牢）",
            priority=SellSignal.PRIORITY["trapped_severe"],
            suggested_action="严重套牢，立即止损",
        )
    elif trapped_level == 4:
        return SellSignal(
            code=position["code"],
            reason="trapped_severe",
            detail=f"套牢盘达到4级（套牢较重）",
            priority=SellSignal.PRIORITY["trapped_severe"],
            suggested_action="套牢加重，考虑减仓",
        )
    
    return None


def check_kline_down(position: Dict, kline_direction: str) -> Optional[SellSignal]:
    """
    检查K线转空
    
    规则：60分钟K线方向为 down 或 strong_down
    """
    if kline_direction in ("down", "strong_down"):
        return SellSignal(
            code=position["code"],
            reason="kline_down",
            detail=f"60分钟K线方向: {kline_direction}",
            priority=SellSignal.PRIORITY["kline_down"],
            suggested_action="K线转空，考虑卖出",
        )
    
    return None


# ============ 持仓检查主流程 ============

def check_positions(
    current_data: Dict[str, Dict] = None,
    market: Dict = None,
    auto_sell: bool = False,
) -> Dict:
    """
    检查所有持仓，生成卖出信号
    
    Args:
        current_data: {code: {price, score, kline, trapped_level, ...}} 当前数据
        market: 大盘分析结果，为None则自动获取
        auto_sell: 是否自动执行卖出（默认False，仅生成信号）
    
    Returns:
        {
            "market_signal": SellSignal or None,
            "position_signals": [{...}, ...],
            "summary": "...",
            "holdings_value": {...},
        }
    """
    logger.info("=" * 50)
    logger.info("持仓健康检查开始")
    
    # 获取大盘数据
    if market is None:
        market = analyze_market()
    
    # 获取持仓
    positions = get_positions("holding")
    if not positions:
        logger.info("当前无持仓")
        return {
            "market_signal": None,
            "position_signals": [],
            "summary": "当前无持仓",
            "holdings_value": None,
        }
    
    logger.info(f"当前持仓: {len(positions)} 只")
    
    # 大盘信号
    market_signal = check_market_condition(market)
    if market_signal:
        logger.warning(f"大盘信号: {market_signal.detail}")
    
    # 获取当前价格（如未提供）
    if current_data is None:
        current_data = {}
        for p in positions:
            code = p["code"]
            try:
                quote = fetch_realtime_quote(code)
                if quote:
                    current_data[code] = {
                        "price": quote.get("price", p["buy_price"]),
                        "score": 0,  # 需要外部传入
                        "kline": "unknown",
                        "trapped_level": 0,
                    }
            except Exception as e:
                logger.warning(f"获取 {code} 行情失败: {e}")
                current_data[code] = {
                    "price": p["buy_price"],
                    "score": 0,
                    "kline": "unknown",
                    "trapped_level": 0,
                }
    
    # 检查每只持仓
    all_signals = []
    
    for position in positions:
        code = position["code"]
        data = current_data.get(code, {})
        current_price = data.get("price", position["buy_price"])
        current_score = data.get("score", 0)
        kline = data.get("kline", "unknown")
        trapped = data.get("trapped_level", 0)
        
        signals = []
        
        # 1. 止损检查
        sig = check_stop_loss(position, current_price)
        if sig:
            signals.append(sig)
        
        # 2. 止盈检查
        sig = check_take_profit(position, current_price)
        if sig:
            signals.append(sig)
        
        # 3. 时间止损
        sig = check_time_stop(position)
        if sig:
            signals.append(sig)
        
        # 4. 评分下降
        sig = check_score_drop(position, current_score)
        if sig:
            signals.append(sig)
        
        # 5. 套牢盘恶化
        sig = check_trapped_level(position, trapped)
        if sig:
            signals.append(sig)
        
        # 6. K线转空
        sig = check_kline_down(position, kline)
        if sig:
            signals.append(sig)
        
        # 按优先级排序
        signals.sort(key=lambda s: s.priority)
        
        if signals:
            top_signal = signals[0]
            logger.warning(f"[{code}] {top_signal.reason}: {top_signal.detail} -> {top_signal.suggested_action}")
            all_signals.append(top_signal.to_dict())
            
            # 自动卖出
            if auto_sell and top_signal.priority <= 5:  # 高优先级信号才自动执行
                sell_position(code, current_price, top_signal.reason)
                logger.info(f"自动卖出: {code} @ {current_price}")
    
    # 计算持仓市值
    prices = {code: data.get("price", 0) for code, data in current_data.items()}
    holdings_value = calculate_position_value(prices)
    
    # 生成摘要
    if all_signals:
        urgent = [s for s in all_signals if s["priority"] <= 3]
        warning = [s for s in all_signals if s["priority"] > 3]
        summary = f"发现 {len(all_signals)} 个信号，其中 {len(urgent)} 个紧急"
    else:
        summary = "持仓健康，无卖出信号"
    
    logger.info(f"检查完成: {summary}")
    logger.info(f"总市值: {holdings_value['total_value']:.0f}，盈亏: {holdings_value['total_profit']:.0f}({holdings_value['total_profit_pct']}%)")
    
    return {
        "market_signal": market_signal.to_dict() if market_signal else None,
        "position_signals": all_signals,
        "summary": summary,
        "holdings_value": holdings_value,
        "market": market,
    }


def generate_position_report(check_result: Dict, output_path: str = None) -> str:
    """
    生成持仓检查报告（docx）
    
    Args:
        check_result: check_positions() 返回的结果
        output_path: 输出路径，默认自动生成
    
    Returns:
        报告文件路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if output_path is None:
        from config.config import OUTPUT_DIR
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(OUTPUT_DIR, f"持仓检查_{ts}.docx")
    
    doc = Document()
    
    # 标题
    heading = doc.add_heading("选股小龙虾 - 持仓健康检查", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 时间
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"检查时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 大盘环境
    market = check_result.get("market", {})
    if market:
        p = doc.add_paragraph()
        p.add_run(f"大盘评分：{market.get('total_score', 0):.0f}分 | {market.get('advice', '')}").italic = True
    
    doc.add_paragraph()
    
    # 持仓概览
    holdings = check_result.get("holdings_value", {})
    if holdings:
        doc.add_heading("持仓概览", level=1)
        doc.add_paragraph(f"持仓数量：{holdings.get('position_count', 0)} 只")
        doc.add_paragraph(f"总成本：{holdings.get('total_cost', 0):.2f}")
        doc.add_paragraph(f"总市值：{holdings.get('total_value', 0):.2f}")
        
        profit = holdings.get('total_profit', 0)
        profit_pct = holdings.get('total_profit_pct', 0)
        p = doc.add_paragraph(f"总盈亏：")
        run = p.add_run(f"{profit:+.2f} ({profit_pct:+.2f}%)")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0) if profit >= 0 else RGBColor(255, 0, 0)
    
    # 卖出信号
    signals = check_result.get("position_signals", [])
    if signals:
        doc.add_heading("⚠️ 卖出信号", level=1)
        
        # 紧急信号
        urgent = [s for s in signals if s["priority"] <= 3]
        if urgent:
            doc.add_heading("🔴 紧急（立即处理）", level=2)
            for sig in urgent:
                p = doc.add_paragraph()
                p.add_run(f"{sig['code']} - {sig['reason']}").bold = True
                doc.add_paragraph(f"详情：{sig['detail']}")
                doc.add_paragraph(f"建议：{sig['suggested_action']}")
                doc.add_paragraph()
        
        # 警告信号
        warning = [s for s in signals if s["priority"] > 3]
        if warning:
            doc.add_heading("🟡 警告（关注）", level=2)
            for sig in warning:
                p = doc.add_paragraph()
                p.add_run(f"{sig['code']} - {sig['reason']}").bold = True
                doc.add_paragraph(f"详情：{sig['detail']}")
                doc.add_paragraph(f"建议：{sig['suggested_action']}")
                doc.add_paragraph()
    else:
        doc.add_heading("✅ 持仓健康", level=1)
        doc.add_paragraph("未发现卖出信号，继续持有")
    
    # 持仓明细
    if holdings and holdings.get("positions"):
        doc.add_heading("持仓明细", level=1)
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        headers = ["代码", "名称", "买入价", "当前价", "盈亏%", "市值", "持有天数"]
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            header_row.cells[i].text = header
            for paragraph in header_row.cells[i].paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].bold = True
        
        for pos in holdings["positions"]:
            row = table.add_row()
            row.cells[0].text = pos["code"]
            row.cells[1].text = pos["name"]
            row.cells[2].text = f"{pos['buy_price']:.2f}"
            row.cells[3].text = f"{pos.get('current_price', pos['buy_price']):.2f}"
            
            profit_pct = pos.get("unrealized_profit_pct", 0)
            cell = row.cells[4]
            cell.text = f"{profit_pct:+.2f}%"
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].font.color.rgb = RGBColor(0, 128, 0) if profit_pct >= 0 else RGBColor(255, 0, 0)
            
            row.cells[5].text = f"{pos.get('market_value', 0):.0f}"
            
            buy_date = datetime.strptime(pos["buy_date"], "%Y-%m-%d")
            hold_days = (datetime.now() - buy_date).days
            row.cells[6].text = str(hold_days)
    
    doc.save(output_path)
    logger.info(f"持仓检查报告已生成: {output_path}")
    return output_path


# ============ CLI 接口 ============

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="持仓健康检查")
    parser.add_argument("--auto-sell", action="store_true", help="自动执行卖出")
    parser.add_argument("--report", action="store_true", help="生成docx报告")
    parser.add_argument("--output", help="报告输出路径")
    
    args = parser.parse_args()
    
    result = check_positions(auto_sell=args.auto_sell)
    
    print(f"\n检查结果: {result['summary']}")
    
    if result["market_signal"]:
        print(f"大盘信号: {result['market_signal']['detail']}")
    
    if result["position_signals"]:
        print(f"\n卖出信号 ({len(result['position_signals'])} 个):")
        for sig in result["position_signals"]:
            print(f"  [{sig['code']}] {sig['reason']}: {sig['detail']}")
    
    if result["holdings_value"]:
        hv = result["holdings_value"]
        print(f"\n持仓概览:")
        print(f"  总市值: {hv['total_value']:.2f}")
        print(f"  总盈亏: {hv['total_profit']:+.2f} ({hv['total_profit_pct']:+.2f}%)")
    
    if args.report:
        report_path = generate_position_report(result, args.output)
        print(f"\n报告已生成: {report_path}")
